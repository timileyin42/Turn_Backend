"""CV/Resume building service for dynamic CV generation, templates, and export functionality."""
from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

from sqlalchemy import and_, desc, func, or_, select, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.template_renderer import render_template
from app.database.cv_models import (
    CV,
    CVExport,
    CVProject,
    CVSection,
    CVSkill,
    CVTemplate,
    Certification,
    Education,
    Language,
    Reference,
    WorkExperience,
)
from app.schemas.cv_schemas import (
    CVAnalyticsResponse,
    CVCreate,
    CVEducationCreate,
    CVEducationResponse,
    CVEducationUpdate,
    CVExperienceCreate,
    CVExperienceResponse,
    CVExperienceUpdate,
    CVExportResponse,
    CVListResponse,
    CVProjectCreate,
    CVProjectResponse,
    CVProjectUpdate,
    CVResponse,
    CVSearchRequest,
    CVSectionCreate,
    CVSectionResponse,
    CVSectionUpdate,
    CVSkillCreate,
    CVSkillResponse,
    CVSkillUpdate,
    CVTemplateResponse,
    CVUpdate,
)

try:  # Optional dependency for PDF generation
    from weasyprint import HTML  # type: ignore[import]
except Exception:  # pragma: no cover - graceful degradation when missing native deps
    HTML = None

try:  # Optional dependency for DOCX exports
    from docx import Document  # type: ignore[import]
except Exception:  # pragma: no cover - allow environments without python-docx
    Document = None


logger = logging.getLogger(__name__)

EXPORT_ROOT = Path(__file__).resolve().parents[2] / "exports"
EXPORT_ROOT.mkdir(parents=True, exist_ok=True)


@dataclass
class _ExportFile:
    """Metadata returned after generating an export file."""
    path: Path
    format: str


def _ensure_list(value: Any) -> List[Any]:
    """Normalize JSON/text columns into list form for templating."""
    if not value:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, list) else [parsed]
        except json.JSONDecodeError:
            return [value]
    return [value]


class CVService:
    """Service for CV/Resume creation, management, and export operations."""
    
    async def create_cv(
        self, 
        db: AsyncSession, 
        user_id: int, 
        cv_data: CVCreate
    ) -> CVResponse:
        """
        Create a new CV for user.
        
        Args:
            db: Database session
            user_id: CV owner ID
            cv_data: CV creation data
            
        Returns:
            Created CV response
        """
        db_cv = CV(
            user_id=user_id,
            **cv_data.model_dump(),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(db_cv)
        await db.commit()
        await db.refresh(db_cv)
        
        return CVResponse.model_validate(db_cv)
    
    async def get_cv_by_id(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int
    ) -> Optional[CVResponse]:
        """
        Get CV by ID with access control.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: Requesting user ID
            
        Returns:
            CV response if user has access
        """
        result = await db.execute(
            select(CV)
            .options(
                selectinload(CV.custom_sections),
                selectinload(CV.educations),
                selectinload(CV.work_experiences),
                selectinload(CV.cv_skills),
                selectinload(CV.projects)
            )
            .where(
                and_(
                    CV.id == cv_id,
                    CV.user_id == user_id
                )
            )
        )
        cv = result.scalar_one_or_none()
        
        if not cv:
            return None
        
        return CVResponse.model_validate(cv)
    
    async def update_cv(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        cv_data: CVUpdate
    ) -> Optional[CVResponse]:
        """
        Update CV information.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID (must be owner)
            cv_data: Updated CV data
            
        Returns:
            Updated CV response
        """
        result = await db.execute(
            select(CV).where(
                and_(
                    CV.id == cv_id,
                    CV.user_id == user_id
                )
            )
        )
        cv = result.scalar_one_or_none()
        
        if not cv:
            return None
        
        # Update fields
        update_data = cv_data.model_dump(exclude_unset=True)
        if update_data:
            for field, value in update_data.items():
                setattr(cv, field, value)
            
            cv.updated_at = datetime.utcnow()
            await db.commit()
            await db.refresh(cv)
        
        return CVResponse.model_validate(cv)
    
    async def delete_cv(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int
    ) -> bool:
        """
        Delete CV (owner only).
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID (must be owner)
            
        Returns:
            True if CV was deleted
        """
        result = await db.execute(
            select(CV).where(
                and_(
                    CV.id == cv_id,
                    CV.user_id == user_id
                )
            )
        )
        cv = result.scalar_one_or_none()
        
        if cv:
            await db.delete(cv)
            await db.commit()
            return True
        
        return False
    
    async def get_user_cvs(
        self, 
        db: AsyncSession, 
        user_id: int,
        skip: int = 0,
        limit: int = 20
    ) -> CVListResponse:
        """
        Get all CVs for a user.
        
        Args:
            db: Database session
            user_id: User ID
            skip: Number of records to skip
            limit: Maximum number of records
            
        Returns:
            Paginated CV list
        """
        # Count total
        count_result = await db.execute(
            select(func.count(CV.id)).where(CV.user_id == user_id)
        )
        total = count_result.scalar()
        
        # Get CVs with pagination
        result = await db.execute(
            select(CV)
            .options(selectinload(CV.custom_sections))
            .where(CV.user_id == user_id)
            .order_by(desc(CV.updated_at))
            .offset(skip)
            .limit(limit)
        )
        cvs = result.scalars().all()
        
        cv_responses = [CVResponse.model_validate(cv) for cv in cvs]
        
        return CVListResponse(
            cvs=cv_responses,
            total=total,
            page=(skip // limit) + 1,
            size=limit,
            pages=(total + limit - 1) // limit
        )
    
    # CV Sections Management
    
    async def create_cv_section(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        section_data: CVSectionCreate
    ) -> Optional[CVSectionResponse]:
        """
        Create a new section in CV.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            section_data: Section creation data
            
        Returns:
            Created section response
        """
        # Check CV ownership
        cv_exists = await self._check_cv_ownership(db, cv_id, user_id)
        if not cv_exists:
            return None
        
        section_payload = section_data.model_dump(exclude={"cv_id"}, exclude_none=True)

        db_section = CVSection(
            cv_id=cv_id,
            **section_payload,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(db_section)
        await db.commit()
        await db.refresh(db_section)
        
        return CVSectionResponse.model_validate(db_section)
    
    async def update_cv_section(
        self, 
        db: AsyncSession, 
        section_id: int, 
        user_id: int, 
        section_data: CVSectionUpdate
    ) -> Optional[CVSectionResponse]:
        """
        Update CV section.
        
        Args:
            db: Database session
            section_id: Section ID
            user_id: User ID
            section_data: Updated section data
            
        Returns:
            Updated section response
        """
        # Get section with CV ownership check
        result = await db.execute(
            select(CVSection)
            .join(CV)
            .where(
                and_(
                    CVSection.id == section_id,
                    CV.user_id == user_id
                )
            )
        )
        section = result.scalar_one_or_none()
        
        if not section:
            return None
        
        # Update fields
        update_data = section_data.model_dump(exclude_unset=True)
        if update_data:
            for field, value in update_data.items():
                setattr(section, field, value)
            
            section.updated_at = datetime.utcnow()
            await db.commit()
            await db.refresh(section)
        
        return CVSectionResponse.model_validate(section)
    
    # CV Education Management
    
    async def add_education(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        education_data: CVEducationCreate
    ) -> Optional[CVEducationResponse]:
        """
        Add education entry to CV.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            education_data: Education data
            
        Returns:
            Created education response
        """
        # Check CV ownership
        if not await self._check_cv_ownership(db, cv_id, user_id):
            return None
        
        education_payload = education_data.model_dump(exclude={"cv_id"}, exclude_none=True)

        db_education = Education(
            cv_id=cv_id,
            **education_payload,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(db_education)
        await db.commit()
        await db.refresh(db_education)
        
        return CVEducationResponse.model_validate(db_education)
    
    async def update_education(
        self, 
        db: AsyncSession, 
        education_id: int, 
        user_id: int, 
        education_data: CVEducationUpdate
    ) -> Optional[CVEducationResponse]:
        """
        Update education entry.
        
        Args:
            db: Database session
            education_id: Education ID
            user_id: User ID
            education_data: Updated education data
            
        Returns:
            Updated education response
        """
        result = await db.execute(
            select(Education)
            .join(CV)
            .where(
                and_(
                    Education.id == education_id,
                    CV.user_id == user_id
                )
            )
        )
        education = result.scalar_one_or_none()
        
        if not education:
            return None
        
        # Update fields
        update_data = education_data.model_dump(exclude_unset=True)
        if update_data:
            for field, value in update_data.items():
                setattr(education, field, value)
            
            education.updated_at = datetime.utcnow()
            await db.commit()
            await db.refresh(education)
        
        return CVEducationResponse.model_validate(education)
    
    # CV Experience Management
    
    async def add_experience(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        experience_data: CVExperienceCreate
    ) -> Optional[CVExperienceResponse]:
        """
        Add work experience entry to CV.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            experience_data: Experience data
            
        Returns:
            Created experience response
        """
        # Check CV ownership
        if not await self._check_cv_ownership(db, cv_id, user_id):
            return None
        
        experience_payload = experience_data.model_dump(exclude={"cv_id"}, exclude_none=True)

        db_experience = WorkExperience(
            cv_id=cv_id,
            **experience_payload,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(db_experience)
        await db.commit()
        await db.refresh(db_experience)
        
        return CVExperienceResponse.model_validate(db_experience)
    
    async def update_experience(
        self, 
        db: AsyncSession, 
        experience_id: int, 
        user_id: int, 
        experience_data: CVExperienceUpdate
    ) -> Optional[CVExperienceResponse]:
        """
        Update experience entry.
        
        Args:
            db: Database session
            experience_id: Experience ID
            user_id: User ID
            experience_data: Updated experience data
            
        Returns:
            Updated experience response
        """
        result = await db.execute(
            select(WorkExperience)
            .join(CV)
            .where(
                and_(
                    WorkExperience.id == experience_id,
                    CV.user_id == user_id
                )
            )
        )
        experience = result.scalar_one_or_none()
        
        if not experience:
            return None
        
        # Update fields
        update_data = experience_data.model_dump(exclude_unset=True)
        if update_data:
            for field, value in update_data.items():
                setattr(experience, field, value)
            
            experience.updated_at = datetime.utcnow()
            await db.commit()
            await db.refresh(experience)
        
        return CVExperienceResponse.model_validate(experience)
    
    # CV Skills Management
    
    async def add_skill(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        skill_data: CVSkillCreate
    ) -> Optional[CVSkillResponse]:
        """
        Add skill to CV.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            skill_data: Skill data
            
        Returns:
            Created skill response
        """
        # Check CV ownership
        if not await self._check_cv_ownership(db, cv_id, user_id):
            return None
        
        skill_payload = skill_data.model_dump(exclude={"cv_id"}, exclude_none=True)

        db_skill = CVSkill(
            cv_id=cv_id,
            **skill_payload,
            created_at=datetime.utcnow()
        )
        
        db.add(db_skill)
        await db.commit()
        await db.refresh(db_skill)
        
        return CVSkillResponse.model_validate(db_skill)
    
    # CV Export functionality
    
    async def export_cv(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int, 
        export_format: str = "pdf",
        template_id: Optional[int] = None
    ) -> Optional[CVExportResponse]:
        """
        Export CV to specified format.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            export_format: Export format (pdf, docx, html)
            template_id: Optional template ID
            
        Returns:
            CV export response with download URL
        """
        # Check CV ownership
        if not await self._check_cv_ownership(db, cv_id, user_id):
            return None
        
        # Fetch CV ORM object with relationships for rendering
        result = await db.execute(
            select(CV)
            .options(
                selectinload(CV.custom_sections),
                selectinload(CV.educations),
                selectinload(CV.work_experiences),
                selectinload(CV.cv_skills),
                selectinload(CV.projects),
                selectinload(CV.certifications),
                selectinload(CV.languages),
                selectinload(CV.references)
            )
            .where(
                and_(
                    CV.id == cv_id,
                    CV.user_id == user_id
                )
            )
        )
        cv_obj = result.scalar_one_or_none()
        if not cv_obj:
            return None
        
        context = self._build_export_context(cv_obj)
        html_content = render_template("cv/modern.html", context)

        export_dir = EXPORT_ROOT / f"user_{user_id}" / f"cv_{cv_id}"
        export_dir.mkdir(parents=True, exist_ok=True)
        base_name = f"cv_{cv_id}_{int(datetime.utcnow().timestamp())}"

        # Create export record
        db_export = CVExport(
            cv_id=cv_id,
            user_id=user_id,
            format=export_format,
            file_url="",
            file_name="",
            file_size=0,
            include_photo=cv_obj.include_photo,
            custom_styling=None,
            created_at=datetime.utcnow()
        )
        
        db.add(db_export)
        await db.commit()
        await db.refresh(db_export)

        html_path = export_dir / f"{base_name}_{db_export.id}.html"
        html_path.write_text(html_content, encoding="utf-8")

        export_file = self._generate_export_file(
            requested_format=export_format,
            html_content=html_content,
            export_dir=export_dir,
            base_name=f"{base_name}_{db_export.id}",
            context=context
        )

        relative_path = export_file.path.relative_to(EXPORT_ROOT).as_posix()
        db_export.format = export_file.format
        db_export.file_url = f"/exports/{relative_path}"
        db_export.file_name = export_file.path.name
        db_export.file_size = export_file.path.stat().st_size
        db_export.expires_at = datetime.utcnow() + timedelta(days=7)

        cv_obj.last_exported_at = datetime.utcnow()
        if export_file.format == "pdf":
            cv_obj.pdf_url = db_export.file_url
        elif export_file.format == "docx":
            cv_obj.docx_url = db_export.file_url

        await db.commit()
        await db.refresh(db_export)

        return CVExportResponse.model_validate(db_export)
    
    async def get_cv_templates(
        self, 
        db: AsyncSession
    ) -> List[CVTemplateResponse]:
        """
        Get available CV templates.
        
        Args:
            db: Database session
            
        Returns:
            List of CV templates
        """
        result = await db.execute(
            select(CVTemplate)
            .where(CVTemplate.is_active == True)
            .order_by(CVTemplate.name)
        )
        templates = result.scalars().all()
        
        return [CVTemplateResponse.model_validate(template) for template in templates]
    
    async def get_cv_analytics(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int
    ) -> Optional[CVAnalyticsResponse]:
        """
        Get CV analytics and statistics.
        
        Args:
            db: Database session
            cv_id: CV ID
            user_id: User ID
            
        Returns:
            CV analytics data
        """
        # Check CV ownership
        if not await self._check_cv_ownership(db, cv_id, user_id):
            return None
        
        # Get CV with all related data
        result = await db.execute(
            select(CV)
            .options(
                selectinload(CV.custom_sections),
                selectinload(CV.educations),
                selectinload(CV.work_experiences),
                selectinload(CV.cv_skills),
                selectinload(CV.projects),
                selectinload(CV.certifications),
                selectinload(CV.languages),
                selectinload(CV.references)
            )
            .where(CV.id == cv_id)
        )
        cv = result.scalar_one_or_none()
        
        if not cv:
            return None
        
        # Calculate completion percentage
        total_sections = 8  # Expected sections: contact, summary, education, experience, skills, projects, achievements, references
        completed_sections = 0
        
        # Basic info completion
        if cv.title:
            completed_sections += 1
        if cv.professional_summary or cv.objective_statement:
            completed_sections += 1
        
        # Section-based completion
        if cv.educations:
            completed_sections += 1
        if cv.work_experiences:
            completed_sections += 1
        if cv.cv_skills:
            completed_sections += 1
        if cv.projects:
            completed_sections += 1
        if cv.certifications:
            completed_sections += 1
        if cv.custom_sections:
            completed_sections += min(len(cv.custom_sections), 2)  # Additional custom sections
        
        completion_percentage = (completed_sections / total_sections) * 100
        
        # Count exports
        export_count_result = await db.execute(
            select(func.count(CVExport.id)).where(CVExport.cv_id == cv_id)
        )
        export_count = export_count_result.scalar()
        
        return {
            "cv_id": cv_id,
            "completion_percentage": completion_percentage,
            "total_sections": len(cv.custom_sections) if cv.custom_sections else 0,
            "education_count": len(cv.educations) if cv.educations else 0,
            "experience_count": len(cv.work_experiences) if cv.work_experiences else 0,
            "skills_count": len(cv.cv_skills) if cv.cv_skills else 0,
            "projects_count": len(cv.projects) if cv.projects else 0,
            "certification_count": len(cv.certifications) if cv.certifications else 0,
            "reference_count": len(cv.references) if cv.references else 0,
            "language_count": len(cv.languages) if cv.languages else 0,
            "export_count": export_count,
            "created_at": cv.created_at,
            "last_updated": cv.updated_at,
        }
    
    # Helper methods
    
    async def _check_cv_ownership(
        self, 
        db: AsyncSession, 
        cv_id: int, 
        user_id: int
    ) -> bool:
        """Check if user owns the CV."""
        result = await db.execute(
            select(CV).where(
                and_(
                    CV.id == cv_id,
                    CV.user_id == user_id
                )
            )
        )
        return result.scalar_one_or_none() is not None

    def _build_export_context(self, cv: CV) -> Dict[str, Any]:
        """Prepare template context for CV exports."""
        experiences = sorted(cv.work_experiences, key=lambda item: item.display_order or 0)
        educations = sorted(cv.educations, key=lambda item: item.display_order or 0)
        skills = sorted(cv.cv_skills, key=lambda item: item.display_order or 0)
        projects = sorted(cv.projects, key=lambda item: item.display_order or 0)
        certifications = sorted(cv.certifications, key=lambda item: item.display_order or 0)
        languages = sorted(cv.languages, key=lambda item: item.display_order or 0)
        references = sorted(cv.references, key=lambda item: item.display_order or 0)
        custom_sections = sorted(cv.custom_sections, key=lambda item: item.order_index or 0)

        return {
            "cv": cv,
            "exported_at": datetime.utcnow(),
            "experiences": experiences,
            "educations": educations,
            "skills": skills,
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
            "references": references,
            "custom_sections": custom_sections,
            "_ensure_list": _ensure_list,
        }

    def _generate_export_file(
        self,
        requested_format: str,
        html_content: str,
        export_dir: Path,
        base_name: str,
        context: Dict[str, Any]
    ) -> _ExportFile:
        """Create the requested export file, falling back to HTML if necessary."""
        requested_format = requested_format.lower()

        if requested_format not in {"pdf", "docx", "html"}:
            logger.warning("Unsupported export format '%s'; falling back to HTML", requested_format)
            requested_format = "html"

        # Always keep the rendered HTML for previews/debugging
        html_path = export_dir / f"{base_name}.html"
        if not html_path.exists():
            html_path.write_text(html_content, encoding="utf-8")

        if requested_format == "pdf" and HTML:
            pdf_path = export_dir / f"{base_name}.pdf"
            try:
                HTML(string=html_content).write_pdf(str(pdf_path))
                return _ExportFile(path=pdf_path, format="pdf")
            except Exception as exc:  # pragma: no cover - depends on native deps
                logger.exception("PDF generation failed; falling back to HTML", exc_info=exc)
        elif requested_format == "docx" and Document:
            docx_path = export_dir / f"{base_name}.docx"
            try:
                self._write_docx(docx_path, context)
                return _ExportFile(path=docx_path, format="docx")
            except Exception as exc:  # pragma: no cover - docx styling variability
                logger.exception("DOCX generation failed; falling back to HTML", exc_info=exc)
        elif requested_format == "pdf" and not HTML:
            logger.warning("WeasyPrint is unavailable; install dependencies to enable PDF exports")
        elif requested_format == "docx" and not Document:
            logger.warning("python-docx is unavailable; install dependencies to enable DOCX exports")

        # Default fallback: return HTML file
        return _ExportFile(path=html_path, format="html")

    def _write_docx(self, docx_path: Path, context: Dict[str, Any]) -> None:
        """Generate a DOCX representation of the CV. Keeps layout simple but readable."""
        if not Document:
            raise RuntimeError("python-docx is not installed")

        document = Document()
        cv = context["cv"]

        document.add_heading(cv.full_name, 0)
        if cv.professional_title:
            document.add_paragraph(cv.professional_title)

        contact_lines = [
            value for value in [cv.email, cv.phone, cv.location, cv.linkedin_url, cv.github_url, cv.portfolio_url]
            if value
        ]
        if contact_lines:
            document.add_paragraph(" | ".join(contact_lines))

        if cv.professional_summary:
            document.add_heading("Summary", level=1)
            document.add_paragraph(cv.professional_summary)

        experiences = context.get("experiences", [])
        if experiences:
            document.add_heading("Experience", level=1)
            for experience in experiences:
                heading = document.add_paragraph()
                heading.add_run(experience.job_title).bold = True
                sub_parts = [part for part in [experience.company_name, experience.location] if part]
                if sub_parts:
                    heading.add_run(f" | {' | '.join(sub_parts)}")

                metadata_parts = []
                if experience.start_date:
                    metadata_parts.append(str(experience.start_date))
                if experience.end_date:
                    metadata_parts.append(str(experience.end_date))
                elif experience.is_current:
                    metadata_parts.append("Present")
                if metadata_parts:
                    document.add_paragraph(" – ".join(metadata_parts))

                if experience.description:
                    document.add_paragraph(experience.description)

                for value in _ensure_list(experience.key_achievements):
                    document.add_paragraph(str(value), style="List Bullet")
                for value in _ensure_list(experience.technologies_used):
                    document.add_paragraph(str(value), style="List Bullet")

        educations = context.get("educations", [])
        if educations:
            document.add_heading("Education", level=1)
            for education in educations:
                heading = document.add_paragraph()
                title = education.degree_type.upper() if education.degree_type else "Education"
                if education.field_of_study:
                    title += f" in {education.field_of_study}"
                heading.add_run(title).bold = True
                sub_parts = [part for part in [education.institution_name, education.location] if part]
                if sub_parts:
                    heading.add_run(f" | {' | '.join(sub_parts)}")

                if education.description:
                    document.add_paragraph(education.description)
                for value in _ensure_list(education.relevant_coursework):
                    document.add_paragraph(str(value), style="List Bullet")
                if education.honors:
                    document.add_paragraph(f"Honors: {education.honors}")

        projects = context.get("projects", [])
        if projects:
            document.add_heading("Projects", level=1)
            for project in projects:
                heading = document.add_paragraph()
                heading.add_run(project.project_name).bold = True
                sub_parts = [part for part in [project.project_type, project.role_in_project] if part]
                if sub_parts:
                    heading.add_run(f" | {' | '.join(sub_parts)}")

                if project.description:
                    document.add_paragraph(project.description)

                for value in _ensure_list(project.key_achievements):
                    document.add_paragraph(str(value), style="List Bullet")
                for value in _ensure_list(project.technologies_used):
                    document.add_paragraph(str(value), style="List Bullet")

        if context.get("skills"):
            document.add_heading("Skills", level=1)
            for skill in context["skills"]:
                skill_line = f"{skill.skill_name} ({skill.skill_category.title()})"
                if skill.proficiency_level:
                    skill_line += f" - {skill.proficiency_level.title()}"
                document.add_paragraph(skill_line, style="List Bullet")

        if context.get("certifications"):
            document.add_heading("Certifications", level=1)
            for certification in context["certifications"]:
                cert_line = certification.certification_name
                if certification.issuing_organization:
                    cert_line += f" - {certification.issuing_organization}"
                document.add_paragraph(cert_line, style="List Bullet")

        if context.get("languages"):
            document.add_heading("Languages", level=1)
            for language in context["languages"]:
                lang_line = f"{language.language_name} ({language.proficiency_level.title()})"
                document.add_paragraph(lang_line, style="List Bullet")

        for section in context.get("custom_sections", []):
            document.add_heading(section.title, level=1)
            document.add_paragraph(section.content)

        document.save(docx_path)


# Global CV service instance
cv_service = CVService()