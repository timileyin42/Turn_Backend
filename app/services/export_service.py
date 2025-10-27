""" 
PDF/DOCX Export Service for CV and Portfolio generation.
Uses free libraries: WeasyPrint (PDF), python-docx (DOCX), ReportLab (PDF)
"""
import os
import io
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime
import base64

# PDF/DOCX libraries
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from jinja2 import Environment, FileSystemLoader
from app.core.config import settings

# Import Cloudinary service for file storage
try:
    from app.services.cloudinary_service import cloudinary_service
    CLOUDINARY_AVAILABLE = True
except ImportError:
    cloudinary_service = None
    CLOUDINARY_AVAILABLE = False


class ExportService:
    """Free export service for generating PDF and DOCX files."""
    
    def __init__(self):
        """Initialize export service with template environment."""
        self.templates_dir = Path(__file__).parent.parent / "templates" / "exports"
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
        # Templates are now stored as separate files in app/templates/exports/
        # cv_template.html and portfolio_template.html

    # PDF Export Methods
    async def export_cv_to_pdf_weasyprint(self, cv_data: Dict[str, Any]) -> bytes:
        """Export CV to PDF using WeasyPrint (free HTML to PDF)."""
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("WeasyPrint not available. Install with: pip install weasyprint")
        
        try:
            template = self.jinja_env.get_template('cv_template.html')
            html_content = template.render(cv=cv_data)
            
            # Convert HTML to PDF
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except Exception as e:
            raise RuntimeError(f"Failed to generate PDF with WeasyPrint: {str(e)}")

    async def export_cv_to_pdf_reportlab(self, cv_data: Dict[str, Any]) -> bytes:
        """Export CV to PDF using ReportLab (free PDF generation)."""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("ReportLab not available. Install with: pip install reportlab")
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []
            
            # Header
            header_style = ParagraphStyle(
                'Header',
                parent=styles['Title'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#007bff'),
                alignment=1  # Center
            )
            story.append(Paragraph(cv_data.get('title', 'CV'), header_style))
            
            # Contact info
            if cv_data.get('user_email') or cv_data.get('user_phone'):
                contact_info = []
                if cv_data.get('user_email'):
                    contact_info.append(cv_data['user_email'])
                if cv_data.get('user_phone'):
                    contact_info.append(cv_data['user_phone'])
                if cv_data.get('user_location'):
                    contact_info.append(cv_data['user_location'])
                
                contact_style = ParagraphStyle(
                    'Contact',
                    parent=styles['Normal'],
                    fontSize=12,
                    spaceAfter=20,
                    alignment=1,
                    textColor=colors.grey
                )
                story.append(Paragraph(' | '.join(contact_info), contact_style))
            
            story.append(Spacer(1, 20))
            
            # Summary
            if cv_data.get('summary'):
                story.append(Paragraph('Professional Summary', styles['Heading2']))
                story.append(Paragraph(cv_data['summary'], styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Experience
            if cv_data.get('experiences'):
                story.append(Paragraph('Work Experience', styles['Heading2']))
                for exp in cv_data['experiences']:
                    story.append(Paragraph(f"<b>{exp.get('job_title', 'N/A')}</b>", styles['Normal']))
                    story.append(Paragraph(f"<i>{exp.get('company_name', 'N/A')}</i>", styles['Normal']))
                    date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    story.append(Paragraph(date_range, styles['Normal']))
                    if exp.get('description'):
                        story.append(Paragraph(exp['description'], styles['Normal']))
                    story.append(Spacer(1, 12))
                story.append(Spacer(1, 20))
            
            # Education
            if cv_data.get('education'):
                story.append(Paragraph('Education', styles['Heading2']))
                for edu in cv_data['education']:
                    degree_info = f"{edu.get('degree', 'N/A')} in {edu.get('field_of_study', 'N/A')}"
                    story.append(Paragraph(f"<b>{degree_info}</b>", styles['Normal']))
                    story.append(Paragraph(f"<i>{edu.get('institution_name', 'N/A')}</i>", styles['Normal']))
                    date_range = f"{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}"
                    story.append(Paragraph(date_range, styles['Normal']))
                    if edu.get('description'):
                        story.append(Paragraph(edu['description'], styles['Normal']))
                    story.append(Spacer(1, 12))
                story.append(Spacer(1, 20))
            
            # Skills
            if cv_data.get('skills'):
                story.append(Paragraph('Skills', styles['Heading2']))
                skills_text = ', '.join([skill.get('skill_name', '') for skill in cv_data['skills']])
                story.append(Paragraph(skills_text, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Projects
            if cv_data.get('projects'):
                story.append(Paragraph('Projects', styles['Heading2']))
                for project in cv_data['projects']:
                    story.append(Paragraph(f"<b>{project.get('project_name', 'N/A')}</b>", styles['Normal']))
                    date_range = f"{project.get('start_date', '')} - {project.get('end_date', 'Present')}"
                    story.append(Paragraph(date_range, styles['Normal']))
                    if project.get('description'):
                        story.append(Paragraph(project['description'], styles['Normal']))
                    if project.get('technologies_used'):
                        tech_text = ', '.join(project['technologies_used'])
                        story.append(Paragraph(f"<i>Technologies: {tech_text}</i>", styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            raise RuntimeError(f"Failed to generate PDF with ReportLab: {str(e)}")

    # DOCX Export Methods
    async def export_cv_to_docx(self, cv_data: Dict[str, Any]) -> bytes:
        """Export CV to DOCX using python-docx (free DOCX generation)."""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Header
            header = doc.add_heading(cv_data.get('title', 'CV'), 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            if cv_data.get('user_email') or cv_data.get('user_phone'):
                contact_info = []
                if cv_data.get('user_email'):
                    contact_info.append(cv_data['user_email'])
                if cv_data.get('user_phone'):
                    contact_info.append(cv_data['user_phone'])
                if cv_data.get('user_location'):
                    contact_info.append(cv_data['user_location'])
                
                contact_para = doc.add_paragraph(' | '.join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            if cv_data.get('summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(cv_data['summary'])
            
            # Experience
            if cv_data.get('experiences'):
                doc.add_heading('Work Experience', level=1)
                for exp in cv_data['experiences']:
                    job_para = doc.add_paragraph()
                    job_para.add_run(exp.get('job_title', 'N/A')).bold = True
                    
                    company_para = doc.add_paragraph()
                    company_para.add_run(exp.get('company_name', 'N/A')).italic = True
                    
                    date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    doc.add_paragraph(date_range)
                    
                    if exp.get('description'):
                        doc.add_paragraph(exp['description'])
                    
                    doc.add_paragraph()  # Space
            
            # Education
            if cv_data.get('education'):
                doc.add_heading('Education', level=1)
                for edu in cv_data['education']:
                    degree_info = f"{edu.get('degree', 'N/A')} in {edu.get('field_of_study', 'N/A')}"
                    degree_para = doc.add_paragraph()
                    degree_para.add_run(degree_info).bold = True
                    
                    institution_para = doc.add_paragraph()
                    institution_para.add_run(edu.get('institution_name', 'N/A')).italic = True
                    
                    date_range = f"{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}"
                    doc.add_paragraph(date_range)
                    
                    if edu.get('description'):
                        doc.add_paragraph(edu['description'])
                    
                    doc.add_paragraph()  # Space
            
            # Skills
            if cv_data.get('skills'):
                doc.add_heading('Skills', level=1)
                skills_text = ', '.join([skill.get('skill_name', '') for skill in cv_data['skills']])
                doc.add_paragraph(skills_text)
            
            # Projects
            if cv_data.get('projects'):
                doc.add_heading('Projects', level=1)
                for project in cv_data['projects']:
                    project_para = doc.add_paragraph()
                    project_para.add_run(project.get('project_name', 'N/A')).bold = True
                    
                    date_range = f"{project.get('start_date', '')} - {project.get('end_date', 'Present')}"
                    doc.add_paragraph(date_range)
                    
                    if project.get('description'):
                        doc.add_paragraph(project['description'])
                    
                    if project.get('technologies_used'):
                        tech_text = ', '.join(project['technologies_used'])
                        tech_para = doc.add_paragraph()
                        tech_para.add_run(f"Technologies: {tech_text}").italic = True
                    
                    doc.add_paragraph()  # Space
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            raise RuntimeError(f"Failed to generate DOCX: {str(e)}")

    # Portfolio Export Methods
    async def export_portfolio_to_pdf(self, portfolio_data: Dict[str, Any]) -> bytes:
        """Export Portfolio to PDF using WeasyPrint."""
        if not WEASYPRINT_AVAILABLE:
            # Fallback to ReportLab if WeasyPrint not available
            return await self.export_portfolio_to_pdf_reportlab(portfolio_data)
        
        try:
            template = self.jinja_env.get_template('portfolio_template.html')
            html_content = template.render(portfolio=portfolio_data)
            
            # Convert HTML to PDF
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except Exception as e:
            raise RuntimeError(f"Failed to generate Portfolio PDF: {str(e)}")

    async def export_portfolio_to_pdf_reportlab(self, portfolio_data: Dict[str, Any]) -> bytes:
        """Export Portfolio to PDF using ReportLab (fallback)."""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("Neither WeasyPrint nor ReportLab available for PDF generation")
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []
            
            # Header with gradient-like styling
            header_style = ParagraphStyle(
                'PortfolioHeader',
                parent=styles['Title'],
                fontSize=28,
                spaceAfter=20,
                textColor=colors.HexColor('#667eea'),
                alignment=1
            )
            story.append(Paragraph(portfolio_data.get('title', 'Portfolio'), header_style))
            
            if portfolio_data.get('description'):
                desc_style = ParagraphStyle(
                    'Description',
                    parent=styles['Normal'],
                    fontSize=14,
                    spaceAfter=30,
                    alignment=1,
                    textColor=colors.grey
                )
                story.append(Paragraph(portfolio_data['description'], desc_style))
            
            story.append(Spacer(1, 20))
            
            # Portfolio Items
            if portfolio_data.get('items'):
                story.append(Paragraph('Portfolio Items', styles['Heading1']))
                for item in portfolio_data['items']:
                    # Item header
                    item_style = ParagraphStyle(
                        'ItemTitle',
                        parent=styles['Heading3'],
                        textColor=colors.HexColor('#333333')
                    )
                    story.append(Paragraph(item.get('title', 'N/A'), item_style))
                    
                    # Item type badge
                    type_style = ParagraphStyle(
                        'ItemType',
                        parent=styles['Normal'],
                        fontSize=10,
                        textColor=colors.white,
                        backColor=colors.HexColor('#667eea'),
                        leftIndent=0,
                        rightIndent=0
                    )
                    story.append(Paragraph(f"  {item.get('item_type', 'N/A').upper()}  ", type_style))
                    
                    # Description
                    if item.get('description'):
                        story.append(Paragraph(item['description'], styles['Normal']))
                    
                    # Technologies
                    if item.get('technologies_used'):
                        tech_text = ', '.join(item['technologies_used'])
                        story.append(Paragraph(f"<i>Technologies: {tech_text}</i>", styles['Normal']))
                    
                    story.append(Spacer(1, 15))
            
            # Achievements
            if portfolio_data.get('achievements'):
                story.append(Paragraph('Achievements', styles['Heading1']))
                for achievement in portfolio_data['achievements']:
                    achievement_style = ParagraphStyle(
                        'AchievementTitle',
                        parent=styles['Heading4'],
                        textColor=colors.HexColor('#28a745')
                    )
                    story.append(Paragraph(achievement.get('title', 'N/A'), achievement_style))
                    if achievement.get('description'):
                        story.append(Paragraph(achievement['description'], styles['Normal']))
                    story.append(Spacer(1, 12))
            
            # Statistics
            story.append(Paragraph('Portfolio Statistics', styles['Heading1']))
            stats_data = [
                ['Portfolio Items', str(len(portfolio_data.get('items', [])))],
                ['Achievements', str(len(portfolio_data.get('achievements', [])))],
                ['Profile Views', str(portfolio_data.get('view_count', 0))]
            ]
            
            stats_table = Table(stats_data, colWidths=[3*inch, 1*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(stats_table)
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            raise RuntimeError(f"Failed to generate Portfolio PDF with ReportLab: {str(e)}")

    # Utility Methods
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats based on installed libraries."""
        formats = []
        
        if WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE:
            formats.append('pdf')
        
        if PYTHON_DOCX_AVAILABLE:
            formats.append('docx')
        
        return formats

    def check_library_availability(self) -> Dict[str, bool]:
        """Check which export libraries are available."""
        return {
            'weasyprint': WEASYPRINT_AVAILABLE,
            'reportlab': REPORTLAB_AVAILABLE,
            'python_docx': PYTHON_DOCX_AVAILABLE
        }

    async def export_cv(self, cv_data: Dict[str, Any], format: str = 'pdf', upload_to_cloud: bool = True) -> Dict[str, Any]:
        """Export CV in specified format and optionally upload to Cloudinary."""
        # Generate the file
        if format.lower() == 'pdf':
            if WEASYPRINT_AVAILABLE:
                file_bytes = await self.export_cv_to_pdf_weasyprint(cv_data)
            elif REPORTLAB_AVAILABLE:
                file_bytes = await self.export_cv_to_pdf_reportlab(cv_data)
            else:
                raise RuntimeError("No PDF generation library available. Install WeasyPrint or ReportLab.")
        
        elif format.lower() == 'docx':
            file_bytes = await self.export_cv_to_docx(cv_data)
        
        else:
            raise ValueError(f"Unsupported format: {format}. Available formats: {self.get_available_formats()}")
        
        result = {
            "file_bytes": file_bytes,
            "format": format,
            "filename": f"cv_{cv_data.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            "size_bytes": len(file_bytes)
        }
        
        # Upload to Cloudinary if requested and available
        if upload_to_cloud and CLOUDINARY_AVAILABLE and cloudinary_service:
            try:
                upload_result = await cloudinary_service.upload_cv_pdf(
                    pdf_content=file_bytes,
                    user_id=cv_data.get('user_id', 0),
                    cv_id=cv_data.get('id', 0)
                )
                
                if upload_result.get('success'):
                    result.update({
                        "cloud_url": upload_result['url'],
                        "public_id": upload_result['public_id'],
                        "cloud_uploaded": True,
                        "presigned_url": cloudinary_service.generate_presigned_url(
                            upload_result['public_id'], 
                            expires_in_hours=24
                        )
                    })
                else:
                    result.update({
                        "cloud_uploaded": False,
                        "cloud_error": upload_result.get('error')
                    })
                    
            except Exception as e:
                result.update({
                    "cloud_uploaded": False,
                    "cloud_error": str(e)
                })
        else:
            result["cloud_uploaded"] = False
            result["cloud_reason"] = "Cloud upload disabled or service unavailable"
        
        return result

    async def export_portfolio(self, portfolio_data: Dict[str, Any], format: str = 'pdf', upload_to_cloud: bool = True) -> Dict[str, Any]:
        """Export Portfolio in specified format and optionally upload to Cloudinary."""
        if format.lower() == 'pdf':
            file_bytes = await self.export_portfolio_to_pdf(portfolio_data)
        else:
            raise ValueError(f"Unsupported format: {format}. Portfolio export currently supports PDF only.")
        
        result = {
            "file_bytes": file_bytes,
            "format": format,
            "filename": f"portfolio_{portfolio_data.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}",
            "size_bytes": len(file_bytes)
        }
        
        # Upload to Cloudinary if requested and available
        if upload_to_cloud and CLOUDINARY_AVAILABLE and cloudinary_service:
            try:
                upload_result = await cloudinary_service.upload_portfolio_pdf(
                    pdf_content=file_bytes,
                    user_id=portfolio_data.get('user_id', 0),
                    portfolio_id=portfolio_data.get('id', 0)
                )
                
                if upload_result.get('success'):
                    result.update({
                        "cloud_url": upload_result['url'],
                        "public_id": upload_result['public_id'],
                        "cloud_uploaded": True,
                        "presigned_url": cloudinary_service.generate_presigned_url(
                            upload_result['public_id'], 
                            expires_in_hours=24
                        )
                    })
                else:
                    result.update({
                        "cloud_uploaded": False,
                        "cloud_error": upload_result.get('error')
                    })
                    
            except Exception as e:
                result.update({
                    "cloud_uploaded": False,
                    "cloud_error": str(e)
                })
        else:
            result["cloud_uploaded"] = False
            result["cloud_reason"] = "Cloud upload disabled or service unavailable"
        
        return result